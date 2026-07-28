"""Safe text extraction for uploaded knowledge-base documents.

The raw file is intentionally kept in memory only and is never persisted.  A
small allow-list plus compressed-container limits keeps this ingestion path
predictable; it is not an arbitrary document converter.
"""

from __future__ import annotations

import io
import json
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARS = 2_000_000
MAX_PDF_PAGES = 200
MAX_ZIP_ENTRIES = 10_000
MAX_ZIP_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_XLSX_SHEETS = 50
MAX_XLSX_ROWS = 100_000
MAX_XLSX_CELLS = 500_000

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx"}
_SOURCE_TYPES = {
    ".txt": "txt",
    ".md": "md",
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
}
_CONTENT_TYPES = {
    ".txt": {"text/plain", "application/octet-stream"},
    ".md": {"text/markdown", "text/plain", "application/octet-stream"},
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
        "application/zip",
    },
    ".xlsx": {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/octet-stream",
        "application/zip",
    },
}


class KnowledgeFileError(ValueError):
    """User-facing upload validation or extraction error."""

    def __init__(self, message: str, *, status_code: int = 422) -> None:
        super().__init__(message)
        self.status_code = status_code


@dataclass(frozen=True, slots=True)
class ExtractedKnowledgeFile:
    title: str
    text: str
    source_type: str


def extract_knowledge_file(
    *,
    filename: str,
    content_type: str | None,
    data: bytes,
    title: str | None = None,
) -> ExtractedKnowledgeFile:
    """Validate one upload and extract bounded plain text from it."""
    safe_name = Path(filename or "").name
    extension = Path(safe_name).suffix.lower()
    is_xlsx = extension == ".xlsx"
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise KnowledgeFileError(
            f"Unsupported file type. Supported extensions: {supported}",
            status_code=415,
        )
    if not data:
        message = "Загруженный файл XLSX пуст" if is_xlsx else "Uploaded file is empty"
        raise KnowledgeFileError(message)
    if len(data) > MAX_UPLOAD_BYTES:
        max_upload_mb = MAX_UPLOAD_BYTES // (1024 * 1024)
        message = (
            f"Размер файла XLSX превышает лимит {max_upload_mb} МБ"
            if is_xlsx
            else f"Uploaded file exceeds the {max_upload_mb} MB limit"
        )
        raise KnowledgeFileError(
            message,
            status_code=413,
        )

    normalized_content_type = (content_type or "application/octet-stream").split(";", 1)[0]
    normalized_content_type = normalized_content_type.strip().lower()
    if normalized_content_type not in _CONTENT_TYPES[extension]:
        message = (
            "Тип содержимого не соответствует файлу XLSX"
            if is_xlsx
            else f"Content type '{normalized_content_type}' does not match {extension}"
        )
        raise KnowledgeFileError(
            message,
            status_code=415,
        )

    stream = io.BytesIO(data)
    if extension in {".txt", ".md"}:
        text = _extract_utf8_text(data)
    elif extension == ".pdf":
        _expect_pdf(data)
        text = _extract_pdf(stream)
    elif extension == ".docx":
        _validate_ooxml(stream, expected_member="word/document.xml", label="DOCX")
        stream.seek(0)
        text = _extract_docx(stream)
    else:
        _validate_ooxml(stream, expected_member="xl/workbook.xml", label="XLSX")
        stream.seek(0)
        text = _extract_xlsx(stream)

    text = _normalize_text(text)
    if not text:
        hint = " OCR is not supported for scanned/image-only files." if extension == ".pdf" else ""
        message = (
            "В файле XLSX не найден текст" if is_xlsx else f"No extractable text was found.{hint}"
        )
        raise KnowledgeFileError(message)
    if len(text) > MAX_EXTRACTED_CHARS:
        message = (
            f"Объём извлечённого текста из XLSX превышает лимит {MAX_EXTRACTED_CHARS:,} символов"
            if is_xlsx
            else f"Extracted text exceeds the {MAX_EXTRACTED_CHARS:,} character limit"
        )
        raise KnowledgeFileError(
            message,
            status_code=413,
        )

    resolved_title = (title or Path(safe_name).stem).strip()
    if not resolved_title:
        message = (
            "Название документа XLSX не может быть пустым"
            if is_xlsx
            else "Document title cannot be empty"
        )
        raise KnowledgeFileError(message)
    if len(resolved_title) > 512:
        message = (
            "Название документа XLSX не может быть длиннее 512 символов"
            if is_xlsx
            else "Document title exceeds 512 characters"
        )
        raise KnowledgeFileError(message)
    return ExtractedKnowledgeFile(
        title=resolved_title,
        text=text,
        source_type=_SOURCE_TYPES[extension],
    )


def parse_upload_tags(raw_tags: str) -> dict[str, str]:
    """Parse a small JSON object supplied as a multipart form field."""
    try:
        value = json.loads(raw_tags or "{}")
    except json.JSONDecodeError as exc:
        raise KnowledgeFileError("Tags must be a valid JSON object") from exc
    if not isinstance(value, dict) or not all(
        isinstance(key, str) and isinstance(item, str) for key, item in value.items()
    ):
        raise KnowledgeFileError("Tags must be a JSON object with string keys and values")
    if len(value) > 50 or any(len(key) > 100 or len(item) > 500 for key, item in value.items()):
        raise KnowledgeFileError("Tags exceed the allowed size")
    return value


def _extract_utf8_text(data: bytes) -> str:
    if b"\x00" in data:
        raise KnowledgeFileError("Text files cannot contain NUL bytes")
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise KnowledgeFileError("TXT and MD files must use UTF-8 encoding") from exc


def _expect_pdf(data: bytes) -> None:
    if not data.lstrip().startswith(b"%PDF-"):
        raise KnowledgeFileError("The uploaded file is not a valid PDF", status_code=415)


def _extract_pdf(stream: BinaryIO) -> str:
    try:
        reader = PdfReader(stream, strict=True)
        if reader.is_encrypted:
            raise KnowledgeFileError("Encrypted PDF files are not supported")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise KnowledgeFileError(f"PDF exceeds the {MAX_PDF_PAGES} page limit", status_code=413)
        parts: list[str] = []
        chars = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            chars += len(page_text)
            if chars > MAX_EXTRACTED_CHARS:
                raise KnowledgeFileError(
                    f"Extracted text exceeds the {MAX_EXTRACTED_CHARS:,} character limit",
                    status_code=413,
                )
            parts.append(page_text)
        return "\n\n".join(parts)
    except KnowledgeFileError:
        raise
    except Exception as exc:
        raise KnowledgeFileError("The PDF could not be read") from exc


def _validate_ooxml(stream: BinaryIO, *, expected_member: str, label: str) -> None:
    is_xlsx = label == "XLSX"
    try:
        with zipfile.ZipFile(stream) as archive:
            members = archive.infolist()
            if len(members) > MAX_ZIP_ENTRIES:
                message = (
                    "Файл XLSX содержит слишком много элементов архива"
                    if is_xlsx
                    else f"{label} contains too many archive entries"
                )
                raise KnowledgeFileError(message, status_code=413)
            unpacked_size = sum(member.file_size for member in members)
            if unpacked_size > MAX_ZIP_UNCOMPRESSED_BYTES:
                max_unpacked_mb = MAX_ZIP_UNCOMPRESSED_BYTES // (1024 * 1024)
                message = (
                    f"Размер распакованного файла XLSX превышает лимит {max_unpacked_mb} МБ"
                    if is_xlsx
                    else f"{label} expands beyond the {max_unpacked_mb} MB limit"
                )
                raise KnowledgeFileError(
                    message,
                    status_code=413,
                )
            names = {member.filename for member in members}
            if "[Content_Types].xml" not in names or expected_member not in names:
                message = (
                    "Загруженный файл не является корректным XLSX"
                    if is_xlsx
                    else f"The uploaded file is not a valid {label}"
                )
                raise KnowledgeFileError(message, status_code=415)
    except KnowledgeFileError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        message = (
            "Загруженный файл не является корректным XLSX"
            if is_xlsx
            else f"The uploaded file is not a valid {label}"
        )
        raise KnowledgeFileError(message, status_code=415) from exc


def _extract_docx(stream: BinaryIO) -> str:
    try:
        document = Document(stream)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [" ".join(cell.text.split()) for cell in row.cells]
                if any(values):
                    parts.append("\t".join(values))
        return _join_bounded(parts)
    except KnowledgeFileError:
        raise
    except Exception as exc:
        raise KnowledgeFileError("The DOCX file could not be read") from exc


def _extract_xlsx(stream: BinaryIO) -> str:
    try:
        workbook = load_workbook(stream, read_only=True, data_only=False, keep_links=False)
        if len(workbook.worksheets) > MAX_XLSX_SHEETS:
            raise KnowledgeFileError(
                f"В файле XLSX больше {MAX_XLSX_SHEETS} листов",
                status_code=413,
            )
        parts: list[str] = []
        rows_seen = 0
        cells_seen = 0
        for sheet in workbook.worksheets:
            max_row = sheet.max_row
            max_column = sheet.max_column
            # Some valid producers omit the optional worksheet dimension.
            # openpyxl then reports unknown bounds in read-only mode but can
            # still iterate safely; the counters below enforce the same limits.
            if (
                max_row is not None
                and max_column is not None
                and (max_row > MAX_XLSX_ROWS or max_row * max_column > MAX_XLSX_CELLS)
            ):
                raise KnowledgeFileError(
                    "Файл XLSX содержит слишком много строк или ячеек",
                    status_code=413,
                )
            sheet_rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                rows_seen += 1
                cells_seen += len(row)
                if rows_seen > MAX_XLSX_ROWS or cells_seen > MAX_XLSX_CELLS:
                    raise KnowledgeFileError(
                        "Файл XLSX содержит слишком много строк или ячеек",
                        status_code=413,
                    )
                values = ["" if value is None else str(value).strip() for value in row]
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    sheet_rows.append("\t".join(values))
            if sheet_rows:
                parts.append(f"Лист: {sheet.title}\n" + "\n".join(sheet_rows))
        workbook.close()
        return _join_bounded(
            parts,
            limit_message=(
                f"Объём извлечённого текста из XLSX превышает лимит "
                f"{MAX_EXTRACTED_CHARS:,} символов"
            ),
        )
    except KnowledgeFileError:
        raise
    except Exception as exc:
        raise KnowledgeFileError("Не удалось прочитать файл XLSX") from exc


def _join_bounded(parts: list[str], *, limit_message: str | None = None) -> str:
    result: list[str] = []
    chars = 0
    for part in parts:
        chars += len(part) + 2
        if chars > MAX_EXTRACTED_CHARS:
            raise KnowledgeFileError(
                limit_message
                or f"Extracted text exceeds the {MAX_EXTRACTED_CHARS:,} character limit",
                status_code=413,
            )
        result.append(part)
    return "\n\n".join(result)


def _normalize_text(text: str) -> str:
    lines = [" ".join(line.replace("\x00", "").split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
